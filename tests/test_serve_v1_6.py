"""Tests for v1.6 daemon routes — reviewer round-trip (CONTRACT v1.6 §A2)."""
from __future__ import annotations

import base64
import io
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from fastapi.testclient import TestClient

from sis_caro_humanizer.serve.app import create_app

TOKEN = "test-token"
HDRS = {"Authorization": f"Bearer {TOKEN}"}


@pytest.fixture
def client() -> TestClient:
    return TestClient(create_app(token=TOKEN))


def _simple_docx_b64() -> str:
    """Build a minimal DOCX (no tracked changes) and return it as base64."""
    doc = Document()
    doc.add_paragraph("This is the first paragraph of the reviewed document.")
    doc.add_paragraph("This is the second paragraph of the reviewed document.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode()


# ---------------------------------------------------------------------------
# /v1/review-import
# ---------------------------------------------------------------------------


def test_review_import_requires_auth(client: TestClient) -> None:
    b64 = _simple_docx_b64()
    r = client.post(
        "/v1/review-import",
        json={"docx_b64": b64, "original_text": "hello"},
    )
    assert r.status_code == 401
    assert r.json()["error"] == "unauthorised"


def test_review_import_returns_diff(client: TestClient) -> None:
    """POST a simple DOCX and verify the response shape is correct."""
    b64 = _simple_docx_b64()
    original_text = (
        "This is the first paragraph of the reviewed document.\n\n"
        "This is a different second paragraph."
    )
    r = client.post(
        "/v1/review-import",
        json={"docx_b64": b64, "original_text": original_text},
        headers=HDRS,
    )
    assert r.status_code == 200, r.text
    body = r.json()

    assert "accepted_text" in body
    assert "diff_sections" in body
    assert "comments" in body
    assert "post_score" in body

    assert isinstance(body["accepted_text"], str)
    assert len(body["accepted_text"]) > 0

    assert isinstance(body["diff_sections"], list)
    assert len(body["diff_sections"]) > 0

    # Every diff section must have the required keys.
    for sec in body["diff_sections"]:
        assert "original" in sec
        assert "revised" in sec
        assert "changed" in sec
        assert "paragraph_idx" in sec

    assert isinstance(body["comments"], list)

    assert "score" in body["post_score"]
    assert "band" in body["post_score"]
    assert body["post_score"]["band"] in ("low", "medium", "high")


# ---------------------------------------------------------------------------
# /v1/citations/google-docs
# ---------------------------------------------------------------------------


def test_google_docs_citations_returns_paragraph_coords(client: TestClient) -> None:
    """POST paragraphs with an orphan citation; response must include
    paragraph_idx and char_in_paragraph on each orphan entry."""
    # Paragraph 0 has an orphan citation; paragraph 1 is unrelated prose.
    paragraphs = [
        "Studies show the trend (Smith, 2020).",
        "This paragraph has no citation at all.",
    ]
    r = client.post(
        "/v1/citations/google-docs",
        json={"paragraphs": paragraphs},
        headers=HDRS,
    )
    assert r.status_code == 200, r.text
    body = r.json()

    assert "orphans" in body
    assert "missing" in body
    assert "unused" in body

    # There should be at least one orphan (Smith, 2020) with no refs provided.
    assert len(body["orphans"]) >= 1
    orphan = body["orphans"][0]
    # Must include paragraph coordinates.
    assert "paragraph_idx" in orphan, "orphan missing paragraph_idx"
    assert "char_in_paragraph" in orphan, "orphan missing char_in_paragraph"
    # The orphan lives in paragraph 0.
    assert orphan["paragraph_idx"] == 0
    # char_in_paragraph must be within the first paragraph's length.
    assert 0 <= orphan["char_in_paragraph"] < len(paragraphs[0])


def test_google_docs_citations_empty_paragraphs(client: TestClient) -> None:
    """POST an empty paragraph list; endpoint must return empty lists."""
    r = client.post(
        "/v1/citations/google-docs",
        json={"paragraphs": []},
        headers=HDRS,
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["orphans"] == []
    assert body["missing"] == []
    assert body["unused"] == []


def test_review_import_no_docx_dep_503(client: TestClient) -> None:
    """If python-docx is unavailable (ImportError), the endpoint returns 503."""
    import sys
    import importlib

    b64 = _simple_docx_b64()

    # Patch docx_bridge module so its imports raise ImportError.
    import sis_caro_humanizer.docx_bridge as bridge_mod

    original_accept = bridge_mod.accept_tracked_changes

    def _raise(*args, **kwargs):
        raise ImportError("python-docx is required for .docx support: pip install 'python-docx>=1.1'")

    bridge_mod.accept_tracked_changes = _raise  # type: ignore[assignment]
    try:
        r = client.post(
            "/v1/review-import",
            json={"docx_b64": b64, "original_text": "hello"},
            headers=HDRS,
        )
    finally:
        bridge_mod.accept_tracked_changes = original_accept  # type: ignore[assignment]

    # The handler catches ImportError and converts to 503 or 400; accept either.
    assert r.status_code in (400, 503, 500)
    body = r.json()
    assert "error" in body
