"""Tests for src/sis_caro_humanizer/docx_bridge.py."""
from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sis_caro_humanizer.docx_bridge import (
    accept_tracked_changes,
    diff_text_sections,
    extract_text,
    extract_word_comments,
    new_docx_from_markdown,
    write_docx,
)


def test_extract_text_roundtrip(tmp_path: Path) -> None:
    """extract_text should return both paragraphs, each present in output."""
    doc = Document()
    doc.add_paragraph("First paragraph here.")
    doc.add_paragraph("Second paragraph here.")
    docx_file = tmp_path / "sample.docx"
    doc.save(str(docx_file))

    result = extract_text(docx_file)

    assert "First paragraph here." in result
    assert "Second paragraph here." in result


def test_write_docx_replaces_text(tmp_path: Path) -> None:
    """write_docx should replace paragraph text with the humanized content."""
    doc = Document()
    doc.add_paragraph("Original first paragraph.")
    doc.add_paragraph("Original second paragraph.")
    original = tmp_path / "original.docx"
    doc.save(str(original))

    humanized = "Humanized first paragraph.\n\nHumanized second paragraph."
    out_path = tmp_path / "humanized.docx"
    write_docx(original, humanized, out_path)

    result_doc = Document(str(out_path))
    texts = [p.text for p in result_doc.paragraphs if p.text.strip()]
    assert texts[0] == "Humanized first paragraph."
    assert texts[1] == "Humanized second paragraph."


def test_new_docx_from_markdown_creates_headings(tmp_path: Path) -> None:
    """new_docx_from_markdown should produce heading paragraphs for # / ## lines."""
    md = "# Introduction\n\nSome body text here.\n\n## Background\n\nMore text."
    out_path = tmp_path / "out.docx"
    new_docx_from_markdown(md, out_path)

    result_doc = Document(str(out_path))
    styles = [p.style.name for p in result_doc.paragraphs if p.text.strip()]
    assert any("Heading 1" in s for s in styles), f"No Heading 1 found in {styles}"
    assert any("Heading 2" in s for s in styles), f"No Heading 2 found in {styles}"
    texts = [p.text for p in result_doc.paragraphs if p.text.strip()]
    assert "Introduction" in texts
    assert "Background" in texts
    assert "Some body text here." in texts


def test_new_docx_from_markdown_bold_italic(tmp_path: Path) -> None:
    """new_docx_from_markdown should create bold and italic runs."""
    md = "This has **bold** and *italic* text."
    out_path = tmp_path / "fmt.docx"
    new_docx_from_markdown(md, out_path)

    result_doc = Document(str(out_path))
    paras = [p for p in result_doc.paragraphs if p.text.strip()]
    assert paras, "Expected at least one paragraph"
    runs = paras[0].runs
    bold_runs = [r for r in runs if r.bold]
    italic_runs = [r for r in runs if r.italic]
    assert bold_runs, "Expected a bold run"
    assert italic_runs, "Expected an italic run"
    assert bold_runs[0].text == "bold"
    assert italic_runs[0].text == "italic"


def test_extract_text_missing_dep(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    """extract_text should raise ImportError when python-docx is not installed."""
    # Create a real .docx so the function reaches the import step.
    doc = Document()
    doc.add_paragraph("Some text.")
    docx_file = tmp_path / "test.docx"
    doc.save(str(docx_file))

    # Simulate python-docx being unavailable by setting sys.modules["docx"] = None.
    # We must also reload the bridge module so _require_docx() re-runs the import.
    import importlib

    import sis_caro_humanizer.docx_bridge as bridge_module

    monkeypatch.setitem(sys.modules, "docx", None)  # type: ignore[arg-type]
    # Reload so the cached import is cleared inside the module.
    monkeypatch.setattr(bridge_module, "_require_docx", lambda: (_ for _ in ()).throw(
        ImportError("python-docx is required for .docx support: pip install 'python-docx>=1.1'")
    ))

    with pytest.raises(ImportError, match="python-docx"):
        bridge_module.extract_text(docx_file)


# ---------------------------------------------------------------------------
# A2.1 — accept_tracked_changes
# ---------------------------------------------------------------------------


def _make_ins_docx(tmp_path: Path, inserted_text: str) -> Path:
    """Create a DOCX with one paragraph containing a w:ins (tracked insertion)."""
    doc = Document()
    # Add a paragraph first so the body has a w:p element to attach to.
    para = doc.add_paragraph("")
    p = para._element

    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "Reviewer")
    ins.set(qn("w:date"), "2026-01-01T00:00:00Z")

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = inserted_text
    r.append(t)
    ins.append(r)
    p.append(ins)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    out = tmp_path / "ins.docx"
    out.write_bytes(buf.read())
    return out


def _make_del_docx(tmp_path: Path, normal_text: str, deleted_text: str) -> Path:
    """Create a DOCX with a normal run followed by a w:del (tracked deletion)."""
    doc = Document()
    # Add a paragraph first so the body has a w:p element.
    para = doc.add_paragraph(normal_text)
    p = para._element

    # Deleted run — should be excluded by accept_tracked_changes.
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), "2")
    del_el.set(qn("w:author"), "Reviewer")
    del_el.set(qn("w:date"), "2026-01-01T00:00:00Z")

    r_del = OxmlElement("w:r")
    dt = OxmlElement("w:delText")
    dt.text = deleted_text
    r_del.append(dt)
    del_el.append(r_del)
    p.append(del_el)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    out = tmp_path / "del.docx"
    out.write_bytes(buf.read())
    return out


def test_accept_tracked_changes_includes_insertions(tmp_path: Path) -> None:
    """Inserted text (w:ins) must appear in the accepted output."""
    docx_path = _make_ins_docx(tmp_path, "inserted text")
    result = accept_tracked_changes(docx_path)
    assert "inserted text" in result


def test_accept_tracked_changes_excludes_deletions(tmp_path: Path) -> None:
    """Deleted text (w:del) must NOT appear; normal run text must be kept."""
    docx_path = _make_del_docx(tmp_path, "keep this", "remove this")
    result = accept_tracked_changes(docx_path)
    assert "keep this" in result
    assert "remove this" not in result


# ---------------------------------------------------------------------------
# A2.1 — extract_word_comments
# ---------------------------------------------------------------------------


def test_extract_word_comments_returns_list(tmp_path: Path) -> None:
    """A document with no comments part should return an empty list."""
    doc = Document()
    doc.add_paragraph("Some academic text here.")
    docx_path = tmp_path / "no_comments.docx"
    doc.save(str(docx_path))

    result = extract_word_comments(docx_path)
    assert isinstance(result, list)
    assert result == []


# ---------------------------------------------------------------------------
# A2.1 — diff_text_sections
# ---------------------------------------------------------------------------


def test_diff_text_sections_marks_changed() -> None:
    """A paragraph that differs between original and revised must have changed=True."""
    original = "First paragraph.\n\nSecond paragraph.\n\nThird paragraph."
    revised = "First paragraph.\n\nModified second paragraph.\n\nThird paragraph."
    sections = diff_text_sections(original, revised)

    changed = [s for s in sections if s["changed"]]
    unchanged = [s for s in sections if not s["changed"]]

    assert len(changed) >= 1, "Expected at least one changed section"
    assert any("Modified" in s["revised"] for s in changed)
    # First and third paragraphs are unchanged.
    assert any("First paragraph." in s["original"] for s in unchanged)
    assert any("Third paragraph." in s["original"] for s in unchanged)


# ---------------------------------------------------------------------------
# B1 — write_docx with citation bookmarks and hyperlinks
# ---------------------------------------------------------------------------


def test_write_docx_adds_reference_bookmarks(tmp_path: Path) -> None:
    """write_docx should inject a w:bookmarkStart with the correct w:name on
    the paragraph that matches the reference entry."""
    # Build an original DOCX with a matching reference paragraph.
    doc = Document()
    doc.add_paragraph("Body text with a citation (Smith, 2020).")
    doc.add_paragraph("Smith, J. (2020). A sample paper. Journal of Tests, 1(1), 1-5.")
    original = tmp_path / "orig.docx"
    doc.save(str(original))

    humanized_text = (
        "Body text with a citation (Smith, 2020).\n\n"
        "## References\n\n"
        "Smith, J. (2020). A sample paper. Journal of Tests, 1(1), 1-5."
    )
    out_path = tmp_path / "out.docx"
    write_docx(original, humanized_text, out_path)

    result_doc = Document(str(out_path))
    # Find a paragraph whose XML contains a bookmarkStart with the expected name.
    found_bookmark = False
    for para in result_doc.paragraphs:
        for bm in para._element.findall(qn("w:bookmarkStart")):
            if bm.get(qn("w:name")) == "ref_smith_2020":
                found_bookmark = True
                break
        if found_bookmark:
            break

    assert found_bookmark, (
        "Expected a w:bookmarkStart with w:name='ref_smith_2020' in the output DOCX"
    )


def test_write_docx_embeds_citation_hyperlinks(tmp_path: Path) -> None:
    """write_docx should convert (Smith, 2020) in prose into a w:hyperlink
    element pointing at ref_smith_2020 when a matching reference bookmark
    exists."""
    doc = Document()
    doc.add_paragraph("Body text with a citation (Smith, 2020).")
    doc.add_paragraph("Smith, J. (2020). A sample paper. Journal of Tests, 1(1), 1-5.")
    original = tmp_path / "orig.docx"
    doc.save(str(original))

    humanized_text = (
        "Body text with a citation (Smith, 2020).\n\n"
        "## References\n\n"
        "Smith, J. (2020). A sample paper. Journal of Tests, 1(1), 1-5."
    )
    out_path = tmp_path / "out.docx"
    write_docx(original, humanized_text, out_path)

    result_doc = Document(str(out_path))
    # Find a paragraph whose XML contains a w:hyperlink with w:anchor=ref_smith_2020.
    found_hyperlink = False
    for para in result_doc.paragraphs:
        for hl in para._element.findall(qn("w:hyperlink")):
            if hl.get(qn("w:anchor")) == "ref_smith_2020":
                found_hyperlink = True
                break
        if found_hyperlink:
            break

    assert found_hyperlink, (
        "Expected a w:hyperlink with w:anchor='ref_smith_2020' in the output DOCX"
    )
