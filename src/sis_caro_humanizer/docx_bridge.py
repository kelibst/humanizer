"""docx_bridge.py — read and write .docx files for the humanizer pipeline.

Public API
----------
extract_text(path)  → str
    Return the full text of a .docx file, paragraphs separated by double
    newlines.

write_docx(original_path, humanized_text, output_path) → None
    Write a new .docx whose paragraph text is replaced by the humanized
    text (preserving the original paragraph styles).  After the replacement
    pass, two additional passes add Word bookmarks on reference entries and
    internal hyperlinks on citation parentheticals in prose.

Both functions raise ``ImportError`` with a helpful message when
``python-docx`` is not installed.
"""
from __future__ import annotations

import re
from pathlib import Path


def _require_docx():
    """Import and return ``docx.Document``, raising ImportError if absent."""
    try:
        from docx import Document  # type: ignore[import]

        return Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for .docx support: "
            "pip install 'python-docx>=1.1'"
        ) from exc


def extract_text(path: Path) -> str:
    """Return the full text of a .docx file, paragraphs separated by ``\\n\\n``."""
    Document = _require_docx()
    doc = Document(str(path))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n\n".join(parts)


def write_docx(original_path: Path, humanized_text: str, output_path: Path) -> None:
    """Write a new .docx at *output_path* with paragraph text replaced by
    *humanized_text*.

    Strategy
    --------
    - Open *original_path* with python-docx.
    - Split *humanized_text* into paragraphs on double-newline.
    - Walk the original document's paragraphs in order.
    - For each original paragraph that has non-empty text:
        - Pop the next humanized paragraph; replace the run text while
          preserving the original paragraph's style.
        - If humanized paragraphs run out, keep the original text.
    - Pass 2: inject ``w:bookmarkStart`` / ``w:bookmarkEnd`` around each
      reference-list paragraph that matches an APA entry from the
      ``## References`` section of *humanized_text*.
    - Pass 3: convert citation parentheticals in prose paragraphs into
      Word internal hyperlinks (``w:hyperlink`` with ``w:anchor``) pointing
      at the bookmarks created in Pass 2.
    - Save to *output_path*.
    """
    Document = _require_docx()
    doc = Document(str(original_path))

    # Skip markdown headings (lines starting with '#') when mapping humanized
    # text back onto DOCX paragraph slots.  Headings exist in markdown but have
    # no direct DOCX paragraph equivalent here; treating them as content would
    # consume an original paragraph slot and offset the rest of the mapping.
    humanized_paras = [
        p.strip()
        for p in humanized_text.split("\n\n")
        if p.strip() and not p.strip().startswith("#")
    ]
    h_iter = iter(humanized_paras)

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        replacement = next(h_iter, None)
        if replacement is None:
            break
        # Clear all runs, then set the text on the first run (preserving style).
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = replacement
        else:
            para.add_run(replacement)

    # Pass 2 — Reference bookmarks
    bookmark_map = _build_reference_bookmarks(doc, humanized_text)

    # Pass 3 — Citation hyperlinks in prose
    _embed_citation_hyperlinks(doc, bookmark_map)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Bookmark helpers (Pass 2)
# ---------------------------------------------------------------------------

# Regex: first line of a typical APA reference (starts with capital, contains
# a 4-digit year in parentheses somewhere on the line).
_APA_LINE_RE = re.compile(r"^[A-Z][a-z].*\(\d{4}\)")

# Regex: extract last name (first word) and year from an APA line.
_APA_KEY_RE = re.compile(r"^([A-Za-z][\w\-']+).*\((\d{4})\)")

# Regex: citation parenthetical in prose, e.g. "(Smith, 2020)" or
# "(Smith & Doe, 2020a)".
_CITE_PAREN_RE = re.compile(
    r"\("
    r"(?P<key>[A-Z][\w\-']+(?:\s+et\s+al\.)?(?:\s*(?:,\s+|\s+(?:and|&)\s+)[A-Z][\w\-']+)*)"
    r"\s*,\s*(?P<year>\d{4}[a-z]?)"
    r"(?:,\s*p{1,2}\.\s*\d+(?:[–\-]\d+)?)?"
    r"\)"
)


def _make_bookmark_id(line: str, existing: dict[str, str]) -> str:
    """Derive a ``ref_lastname_year`` bookmark id from an APA reference line.

    Handles collisions by appending ``_a``, ``_b``, … suffixes.
    """
    m = _APA_KEY_RE.match(line.strip())
    if not m:
        return ""
    last = m.group(1).lower()
    year = m.group(2)
    base = f"ref_{last}_{year}"
    if base not in existing.values():
        return base
    # Collision handling — try _a, _b, … _z suffixes.
    for c in "abcdefghijklmnopqrstuvwxyz":
        candidate = f"{base}_{c}"
        if candidate not in existing.values():
            return candidate
    return base  # give up; caller may overwrite


def _inject_bookmark(paragraph: object, bookmark_id: str, counter: int) -> None:
    """Wrap the first run of *paragraph* with a Word bookmark.

    The bookmark id attribute (``w:id``) must be a non-negative integer.
    We use *counter* (starting at 1000) as the numeric id.

    Parameters
    ----------
    paragraph:
        A ``python-docx`` ``Paragraph`` object.
    bookmark_id:
        The symbolic name for the bookmark (e.g. ``ref_smith_2020``).
    counter:
        Numeric id for this bookmark (must be unique within the document).
    """
    from docx.oxml.ns import qn as _qn  # type: ignore[import]
    from docx.oxml import OxmlElement as _OxmlElement  # type: ignore[import]

    bm_start = _OxmlElement("w:bookmarkStart")
    bm_start.set(_qn("w:id"), str(counter))
    bm_start.set(_qn("w:name"), bookmark_id)

    bm_end = _OxmlElement("w:bookmarkEnd")
    bm_end.set(_qn("w:id"), str(counter))

    runs = paragraph._element.findall(_qn("w:r"))  # type: ignore[attr-defined]
    if runs:
        runs[0].addprevious(bm_start)
        runs[-1].addnext(bm_end)
    else:
        # No runs yet — append at the tail of the paragraph element.
        paragraph._element.append(bm_start)  # type: ignore[attr-defined]
        paragraph._element.append(bm_end)  # type: ignore[attr-defined]


def _build_reference_bookmarks(doc: object, humanized_text: str) -> dict[str, str]:
    """Scan the ``## References`` section of *humanized_text* and inject Word
    bookmarks into the matching DOCX paragraphs.

    Returns a mapping ``{bookmark_id: normalized_para_text}`` for all bookmarks
    that were successfully injected (used in Pass 3).
    """
    # Extract reference lines from the humanized text.
    ref_lines: list[str] = []
    in_refs = False
    for raw_line in humanized_text.splitlines():
        line = raw_line.strip()
        if not in_refs:
            if line.startswith("## References"):
                in_refs = True
            continue
        # Stop at the next heading.
        if line.startswith("#"):
            break
        if line and not line.startswith("#"):
            ref_lines.append(line)

    # Build a mapping: normalised line → bookmark id.
    line_to_bm: dict[str, str] = {}
    for line in ref_lines:
        if not _APA_LINE_RE.match(line):
            continue
        bm_id = _make_bookmark_id(line, {v: v for v in line_to_bm.values()})
        if bm_id:
            line_to_bm[line] = bm_id

    if not line_to_bm:
        return {}

    # Inject bookmarks into matching DOCX paragraphs.
    bookmark_map: dict[str, str] = {}  # bookmark_id → para text
    counter = 1000
    for para in doc.paragraphs:  # type: ignore[attr-defined]
        para_text = para.text.strip()
        if para_text in line_to_bm:
            bm_id = line_to_bm[para_text]
            _inject_bookmark(para, bm_id, counter)
            bookmark_map[bm_id] = para_text
            counter += 1

    return bookmark_map


# ---------------------------------------------------------------------------
# Hyperlink helpers (Pass 3)
# ---------------------------------------------------------------------------


def _add_internal_hyperlink(paragraph: object, anchor: str, match_text: str) -> None:
    """Replace the portion of *paragraph*'s run text that equals *match_text*
    with a ``w:hyperlink`` element pointing at *anchor*.

    Because paragraph runs may be split arbitrarily by python-docx we use a
    destructive-rebuild strategy:

    1. Find the full paragraph text and locate *match_text* within it.
    2. Clear all runs.
    3. Re-add three synthetic runs: prefix text, hyperlink run, suffix text.
    """
    from docx.oxml.ns import qn as _qn  # type: ignore[import]
    from docx.oxml import OxmlElement as _OxmlElement  # type: ignore[import]

    full_text = paragraph.text  # type: ignore[attr-defined]
    idx = full_text.find(match_text)
    if idx == -1:
        return

    # Capture style from first run (if any) for font preservation.
    old_runs = paragraph.runs  # type: ignore[attr-defined]
    rPr_xml: str | None = None
    if old_runs:
        rpr_el = old_runs[0]._element.find(_qn("w:rPr"))
        if rpr_el is not None:
            import xml.etree.ElementTree as _ET
            rPr_xml = _ET.tostring(rpr_el, encoding="unicode")

    # Clear all existing runs from the paragraph element.
    p_el = paragraph._element  # type: ignore[attr-defined]
    for r_el in list(p_el.findall(_qn("w:r"))):
        p_el.remove(r_el)
    for hl_el in list(p_el.findall(_qn("w:hyperlink"))):
        p_el.remove(hl_el)

    def _make_run(text: str) -> object:
        r = _OxmlElement("w:r")
        if rPr_xml:
            import xml.etree.ElementTree as _ET
            rpr_copy = _ET.fromstring(rPr_xml)
            r.append(rpr_copy)
        t = _OxmlElement("w:t")
        t.text = text
        if text and (text[0] == " " or text[-1] == " "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        return r

    prefix = full_text[:idx]
    suffix = full_text[idx + len(match_text):]

    if prefix:
        p_el.append(_make_run(prefix))

    # Build the hyperlink element.
    hyperlink = _OxmlElement("w:hyperlink")
    hyperlink.set(_qn("w:anchor"), anchor)
    link_run = _OxmlElement("w:r")
    if rPr_xml:
        import xml.etree.ElementTree as _ET
        rpr_copy = _ET.fromstring(rPr_xml)
        link_run.append(rpr_copy)
    link_t = _OxmlElement("w:t")
    link_t.text = match_text
    if match_text and (match_text[0] == " " or match_text[-1] == " "):
        link_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    link_run.append(link_t)
    hyperlink.append(link_run)
    p_el.append(hyperlink)

    if suffix:
        p_el.append(_make_run(suffix))


def _cite_anchor(cite_key: str, cite_year: str) -> str:
    """Derive the bookmark id from a citation key and year string."""
    # Strip trailing letter from year (2020a → 2020).
    year_digits = re.match(r"(\d{4})", cite_year)
    year = year_digits.group(1) if year_digits else cite_year
    # First author last name.
    key_clean = re.sub(r"\s*&\s*", " ", cite_key)
    key_clean = re.sub(r"\s+and\s+", " ", key_clean, flags=re.IGNORECASE)
    key_clean = re.sub(r"\s+et\s+al\.?", "", key_clean, flags=re.IGNORECASE)
    m = re.match(r"[A-Za-z][\w\-']+", key_clean)
    if not m:
        return ""
    last = m.group(0).lower()
    return f"ref_{last}_{year}"


def _embed_citation_hyperlinks(doc: object, bookmark_map: dict[str, str]) -> None:
    """Scan prose paragraphs and convert citation parentheticals to hyperlinks.

    Only paragraphs whose text does NOT look like a reference entry are
    processed (we skip lines that themselves are APA references).
    """
    if not bookmark_map:
        return

    for para in doc.paragraphs:  # type: ignore[attr-defined]
        para_text = para.text
        if not para_text.strip():
            continue
        # Skip reference-list paragraphs.
        if _APA_LINE_RE.match(para_text.strip()):
            continue

        # Find all citation parentheticals in this paragraph, left-to-right.
        matches = list(_CITE_PAREN_RE.finditer(para_text))
        if not matches:
            continue

        # Process citations left-to-right.  After each rebuild the para text
        # changes, so we re-scan from scratch.  We track which anchors we have
        # already linked so we don't loop infinitely.
        linked: set[str] = set()
        for _ in range(len(matches)):
            current_text = para.text
            m = _CITE_PAREN_RE.search(current_text)
            if not m:
                break
            anchor = _cite_anchor(m.group("key"), m.group("year"))
            match_text = m.group(0)
            if anchor not in bookmark_map or anchor in linked:
                # Try the next citation if this one cannot be linked.
                # Since we are scanning left-to-right and only skip unresolvable
                # ones, we must stop here to avoid an infinite loop.
                break
            _add_internal_hyperlink(para, anchor, match_text)
            linked.add(anchor)


def accept_tracked_changes(path: Path) -> str:
    """Return the document text after accepting all tracked changes.

    - Text inside ``w:ins`` elements (insertions) is included.
    - Text inside ``w:del`` elements (deletions) is excluded entirely.
    - Normal run text is included unchanged.

    Paragraphs are joined with ``\\n\\n``.  Empty paragraphs are dropped.
    """
    Document = _require_docx()
    doc = Document(str(path))

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    paragraphs_text: list[str] = []
    for p_el in doc._element.findall(f".//{{{WNS}}}p"):
        parts: list[str] = []

        for child in p_el.iter():
            tag_local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            # Skip the text nodes that are inside w:del
            if tag_local == "del":
                # We skip the entire subtree by recording nothing; iter() visits
                # children too, but we only collect w:t text below, so we need
                # to mark deleted runs so their w:t nodes are ignored.
                # The cleanest approach: collect from direct w:r and w:ins children.
                pass

        # Re-walk collecting text correctly: normal runs and inserted runs only.
        for child in p_el:
            tag_local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag_local == "del":
                # Skip deleted runs entirely.
                continue
            elif tag_local == "ins":
                # Accept insertions: include w:r/w:t inside.
                for t_el in child.findall(f".//{{{WNS}}}t"):
                    if t_el.text:
                        parts.append(t_el.text)
            elif tag_local == "r":
                # Normal run.
                for t_el in child.findall(f".//{{{WNS}}}t"):
                    if t_el.text:
                        parts.append(t_el.text)

        text = "".join(parts).strip()
        if text:
            paragraphs_text.append(text)

    return "\n\n".join(paragraphs_text)


def extract_word_comments(path: Path) -> list[dict]:
    """Extract reviewer comments from a .docx file.

    Returns a list of dicts with keys:
      ``id``, ``author``, ``date``, ``text``, ``paragraph_idx``

    ``paragraph_idx`` is the approximate index (0-based) of the paragraph in
    the main body that anchors the comment.  When no anchor is found it is -1.

    Returns ``[]`` when the document has no comments part.
    """
    Document = _require_docx()
    doc = Document(str(path))

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    REL_COMMENTS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

    # Locate the comments part via the relationship table.
    try:
        comments_part = doc.part.part_related_by(REL_COMMENTS)
    except KeyError:
        return []

    comments_el = comments_part._element
    results: list[dict] = []

    # Build a map from comment-id → paragraph index using the main body.
    # Walk all w:p elements in document order and note which carry a
    # w:commentReference with each id.
    body_paragraphs = doc._element.findall(f".//{{{WNS}}}p")
    id_to_para_idx: dict[str, int] = {}
    for para_idx, p_el in enumerate(body_paragraphs):
        for ref_el in p_el.findall(f".//{{{WNS}}}commentReference"):
            ref_id = ref_el.get(f"{{{WNS}}}id")
            if ref_id is not None and ref_id not in id_to_para_idx:
                id_to_para_idx[ref_id] = para_idx

    # Parse each w:comment element.
    for comment_el in comments_el.findall(f".//{{{WNS}}}comment"):
        c_id = comment_el.get(f"{{{WNS}}}id", "")
        author = comment_el.get(f"{{{WNS}}}author", "")
        date = comment_el.get(f"{{{WNS}}}date", "")
        # Join all w:t text nodes inside the comment.
        text_parts = [
            t.text for t in comment_el.findall(f".//{{{WNS}}}t") if t.text
        ]
        text = "".join(text_parts)
        para_idx = id_to_para_idx.get(c_id, -1)
        results.append(
            {
                "id": c_id,
                "author": author,
                "date": date,
                "text": text,
                "paragraph_idx": para_idx,
            }
        )

    return results


def diff_text_sections(original: str, revised: str) -> list[dict]:
    """Compare two texts paragraph-by-paragraph and return a diff report.

    Parameters
    ----------
    original:
        The original text (before lecturer changes).
    revised:
        The revised / accepted text (after ``accept_tracked_changes``).

    Returns
    -------
    list of dicts with keys:
      ``original`` (str) — original paragraph text (empty string if inserted)
      ``revised``  (str) — revised paragraph text (empty string if deleted)
      ``changed``  (bool)
      ``paragraph_idx`` (int) — 0-based index in the *revised* list
    """
    import difflib

    def _split(text: str) -> list[str]:
        return [p.strip() for p in text.split("\n\n") if p.strip()]

    orig_paras = _split(original)
    rev_paras = _split(revised)

    matcher = difflib.SequenceMatcher(None, orig_paras, rev_paras, autojunk=False)
    sections: list[dict] = []

    rev_idx = 0  # tracks position in the revised list

    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        if opcode == "equal":
            for k, para in enumerate(rev_paras[j1:j2]):
                sections.append(
                    {
                        "original": para,
                        "revised": para,
                        "changed": False,
                        "paragraph_idx": j1 + k,
                    }
                )
        elif opcode == "replace":
            # Pair up as many orig ↔ revised as possible; extras are inserts/deletes.
            orig_chunk = orig_paras[i1:i2]
            rev_chunk = rev_paras[j1:j2]
            max_len = max(len(orig_chunk), len(rev_chunk))
            for k in range(max_len):
                orig_text = orig_chunk[k] if k < len(orig_chunk) else ""
                rev_text = rev_chunk[k] if k < len(rev_chunk) else ""
                # paragraph_idx in the revised list — use j1+k for revised side
                p_idx = j1 + k if k < len(rev_chunk) else j2 - 1
                sections.append(
                    {
                        "original": orig_text,
                        "revised": rev_text,
                        "changed": True,
                        "paragraph_idx": p_idx,
                    }
                )
        elif opcode == "delete":
            for k, para in enumerate(orig_paras[i1:i2]):
                sections.append(
                    {
                        "original": para,
                        "revised": "",
                        "changed": True,
                        "paragraph_idx": j1,  # insertion point in revised
                    }
                )
        elif opcode == "insert":
            for k, para in enumerate(rev_paras[j1:j2]):
                sections.append(
                    {
                        "original": "",
                        "revised": para,
                        "changed": True,
                        "paragraph_idx": j1 + k,
                    }
                )

    return sections


__all__ = [
    "extract_text",
    "write_docx",
    "accept_tracked_changes",
    "extract_word_comments",
    "diff_text_sections",
    # B1 helpers (exported for tests)
    "_make_bookmark_id",
    "_inject_bookmark",
    "_add_internal_hyperlink",
]
